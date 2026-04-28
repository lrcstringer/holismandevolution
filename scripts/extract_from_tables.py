"""
Extract chapter content from two-column Word tables.
Source: Desktop/chapters/chapter_N_moderni*.docx

Each document has one table:
  - Column 1: original text  (summary rows are italic)
  - Column 2: modernised text
  - Each row = one paragraph (already manually aligned)

Rules applied:
  - Summary paragraphs (italic col 1): output as plain text (no italics marker)
  - First non-summary paragraph: first word forced to ALL CAPS in both columns

Outputs:
  src/content/original/chapter-XX.md
  src/content/modernised/chapter-XX.md

Run from project root: python scripts/extract_from_tables.py
"""

import os
import re
import glob
from docx import Document

CHAPTERS_DIR = r"C:\Users\lance\OneDrive\Desktop\chapters"
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_ORIGINAL = os.path.join(PROJECT_ROOT, "src", "content", "original")
OUT_MODERNISED = os.path.join(PROJECT_ROOT, "src", "content", "modernised")

CHAPTER_TITLES = [
    "The Reform of Fundamental Concepts",
    "The Reformed Concepts of Space and Time",
    "The Reformed Concept of Matter",
    "The Cell and the Organism",
    "General Concept of Holism",
    "Some Holistic Functions and Categories",
    "Mechanism and Holism",
    "Darwinism and Holism",
    "Mind as an Organ of Wholes",
    "Personality as a Whole",
    "Some Functions and Ideals of Personality",
    "The Holistic Universe",
]


def is_italic_cell(cell) -> bool:
    """Return True if every non-empty run in the cell's paragraphs is italic."""
    runs = [r for p in cell.paragraphs for r in p.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.italic for r in runs)


def cell_text(cell) -> str:
    """Plain text from a cell, stripping leading/trailing whitespace."""
    return cell.text.strip()


def first_word_caps(text: str) -> str:
    """Force the first word of text to ALL CAPS."""
    m = re.match(r"^(\s*)(\S+)(.*)", text, re.DOTALL)
    if m:
        return m.group(1) + m.group(2).upper() + m.group(3)
    return text


def find_chapter_files() -> dict[int, str]:
    """Return {chapter_num: filepath} for all docx files in CHAPTERS_DIR."""
    mapping = {}
    for path in glob.glob(os.path.join(CHAPTERS_DIR, "*.docx")):
        name = os.path.basename(path)
        m = re.search(r"chapter_(\d+)_moderni", name, re.IGNORECASE)
        if m:
            mapping[int(m.group(1))] = path
    return mapping


def process_chapter(chapter_num: int, docx_path: str) -> None:
    title = CHAPTER_TITLES[chapter_num - 1]
    doc = Document(docx_path)

    if not doc.tables:
        print(f"  Ch {chapter_num}: WARNING — no table found, skipping")
        return

    table = doc.tables[0]
    orig_blocks: list[str] = []
    mod_blocks: list[str] = []
    found_body = False  # True once we've left the summary

    for row in table.rows:
        if len(row.cells) < 2:
            continue

        orig_cell = row.cells[0]
        mod_cell = row.cells[1]
        orig_text = cell_text(orig_cell)
        mod_text = cell_text(mod_cell)

        if not orig_text and not mod_text:
            continue

        summary_row = is_italic_cell(orig_cell)

        if not summary_row and not found_body:
            # First paragraph of the chapter proper
            found_body = True
            orig_text = first_word_caps(orig_text)
            mod_text = first_word_caps(mod_text)

        if orig_text:
            orig_blocks.append(orig_text)
        if mod_text:
            mod_blocks.append(mod_text)

    # Write original
    orig_body = "\n\n".join(orig_blocks)
    orig_md = f'---\ntitle: "{title}"\nchapter: {chapter_num}\n---\n\n{orig_body}\n'
    orig_path = os.path.join(OUT_ORIGINAL, f"chapter-{chapter_num:02d}.md")
    with open(orig_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(orig_md)

    # Write modernised
    mod_body = "\n\n".join(mod_blocks)
    mod_md = f'---\ntitle: "{title}"\nchapter: {chapter_num}\n---\n\n{mod_body}\n'
    mod_path = os.path.join(OUT_MODERNISED, f"chapter-{chapter_num:02d}.md")
    with open(mod_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(mod_md)

    print(
        f"  Ch {chapter_num:2d} -- {len(orig_blocks)} rows "
        f"(orig) / {len(mod_blocks)} rows (mod)"
    )


def main():
    os.makedirs(OUT_ORIGINAL, exist_ok=True)
    os.makedirs(OUT_MODERNISED, exist_ok=True)

    chapter_files = find_chapter_files()
    if not chapter_files:
        print(f"No chapter files found in: {CHAPTERS_DIR}")
        return

    for num in sorted(chapter_files):
        process_chapter(num, chapter_files[num])

    missing = [i for i in range(1, 13) if i not in chapter_files]
    if missing:
        print(f"\nMissing chapters: {missing}")
    else:
        print("\nAll 12 chapters extracted.")


if __name__ == "__main__":
    main()
