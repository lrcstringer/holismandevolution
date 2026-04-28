"""
Export chapter markdown files to Word docx.
Outputs to docx-export/ in the project root.
Run from the project root: python scripts/export_docx.py
"""

import os
import re
from docx import Document
from docx.shared import Pt

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONTENT_DIR = os.path.join(PROJECT_ROOT, "src", "content")
OUT_DIR = os.path.join(PROJECT_ROOT, "docx-export")

CHAPTER_TITLES = [
    "The Reform of Fundamental Concepts",
    "The Reformed Concepts of Space and Time",
    "The Reformed Concept of Matter",
    "The Cell and the Organism",
    "General Concept of Holism",
    "Some Holistic Functions and Categories",
    "Mechanism and Holism",
    "Darwinism and Holism",
    "Mind as an Organ of Wholes",
    "Personality as a Whole",
    "Some Functions and Ideals of Personality",
    "The Holistic Universe",
]


def strip_frontmatter(text: str) -> str:
    text = text.strip()
    if text.startswith("---"):
        end = text.find("---", 3)
        if end != -1:
            return text[end + 3:].strip()
    return text


def apply_inlines(para, text: str) -> None:
    """Write text into a paragraph run, respecting **bold** and *italic*."""
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    for part in pattern.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            para.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            para.add_run(part[1:-1]).italic = True
        else:
            para.add_run(part)


def build_blocks(md_body: str) -> list[str]:
    """Split body into non-empty blocks separated by blank lines."""
    blocks = []
    current: list[str] = []
    for line in md_body.splitlines():
        if line.strip() == "":
            if current:
                blocks.append("\n".join(current))
                current = []
        else:
            current.append(line)
    if current:
        blocks.append("\n".join(current))
    return [b.strip() for b in blocks if b.strip()]


def md_to_docx(md_text: str, chapter_num: int, title: str, version: str) -> Document:
    doc = Document()

    # Document title
    heading = doc.add_heading(level=0)
    heading.add_run(f"Chapter {chapter_num}").bold = True
    heading.add_run(f" — {title}")

    version_label = "Original Text (1926)" if version == "original" else "Modernised Text (Centennial Edition)"
    sub = doc.add_paragraph()
    sub.add_run(version_label).italic = True
    sub.paragraph_format.space_after = Pt(18)

    body = strip_frontmatter(md_text)

    for block in build_blocks(body):
        # Headings
        if block.startswith("#### "):
            doc.add_heading(block[5:], level=4)
        elif block.startswith("### "):
            doc.add_heading(block[4:], level=3)
        elif block.startswith("## "):
            doc.add_heading(block[3:], level=2)
        elif block.startswith("# "):
            doc.add_heading(block[2:], level=1)

        # Blockquote (Summary)
        elif block.startswith("> "):
            content = block[2:]
            # Strip markdown bold from Summary label if present
            content = re.sub(r"\*\*Summary:\*\*\s*", "", content)
            # Remaining lines may also start with '> '
            content = re.sub(r"\n> ?", " ", content)
            para = doc.add_paragraph(style="Quote")
            label_run = para.add_run("Summary: ")
            label_run.bold = True
            apply_inlines(para, content)

        # Normal paragraph
        else:
            # Join continuation lines (multi-line block = one paragraph)
            text = re.sub(r"\n", " ", block)
            para = doc.add_paragraph(style="Normal")
            apply_inlines(para, text)

    return doc


def main():
    os.makedirs(OUT_DIR, exist_ok=True)

    for i in range(1, 13):
        title = CHAPTER_TITLES[i - 1]
        ch_str = f"{i:02d}"

        for version in ("original", "modernised"):
            md_path = os.path.join(CONTENT_DIR, version, f"chapter-{ch_str}.md")
            if not os.path.exists(md_path):
                print(f"MISSING: {md_path}")
                continue

            with open(md_path, encoding="utf-8") as f:
                md_text = f.read()

            doc = md_to_docx(md_text, i, title, version)

            out_name = f"chapter-{ch_str}-{version}.docx"
            out_path = os.path.join(OUT_DIR, out_name)
            doc.save(out_path)
            print(f"  {out_name}")

    print(f"\nDone. Files written to:\n  {OUT_DIR}")


if __name__ == "__main__":
    main()
